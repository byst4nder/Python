import docx
from docx.text.font import Font
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.shared import RGBColor


'''
从txt读取文本，生成word文件，并调整格式。
'''

# 读取文本文件，按照行读取
def txtRead(filename):
    with open(filename, 'r', encoding='utf-8') as t:
        for txtline in t:
            yield txtline
        yield None



# 将内容写入word文档
def saveDocx(docFilename,text:str):
    doc = docx.Document(docFilename)
    # font = Font
    # font.name = '仿宋'
    # font.size = Pt(16)
    # font.color.rgb = RGBColor(0x42, 0x24, 0xE9)


    para = doc.add_paragraph(text)
    # para.style.font.name = '仿宋'
    para.style.font.size = Pt(16)
    # 设置段落字体加粗
    # para.style.font = font

    if text.startswith('一'):
        print('一级标题',text)
        para.style._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    else:
        para.style._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
        # para.style.font.size = Pt(12)
    # 保存word文件
    doc.save(docFilename)


# 设置文本格式 判断序号判断标题、设置格式



if __name__ == '__main__':
    filename = 'demo.txt'
    docFilename = 'demo.docx'
    txt = txtRead(filename)
    doc = docx.Document()

    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'),'仿宋')

    doc.save(docFilename)
    while True:
        txtLine = txt.__next__()
        if txtLine:
            txtLine = txtLine.replace('\n', '')
            if txtLine:
                saveDocx(docFilename,txtLine)
                print(txtLine)
        else:
            break


